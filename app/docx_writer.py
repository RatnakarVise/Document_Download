# app/docx_writer.py
from docx import Document
from docx.shared import Pt
from io import BytesIO

def create_technical_spec_docx(text: str, file_obj: BytesIO):
    """
    Creates a technical specification DOCX file in the given file object.
    """
    doc = Document()
    doc.add_heading('Technical Specification', level=1)
    para = doc.add_paragraph(text)
    para.style.font.size = Pt(11)
    doc.save(file_obj)

def create_abap_code_docx(abap_code: str, file_obj: BytesIO):
    """
    Creates an ABAP code DOCX file in the given file object, using monospaced font.
    """
    doc = Document()
    doc.add_heading('ABAP Code', level=1)
    para = doc.add_paragraph()
    run = para.add_run(abap_code)
    run.font.name = 'Consolas'  # Monospaced font for code
    run.font.size = Pt(10.5)
    doc.save(file_obj)

def create_functional_spec_docx(text: str, file_obj: BytesIO):
    """
    Creates a functional specification DOCX file in the given file object.
    """
    doc = Document()
    doc.add_heading('Functional Specification', level=1)
    para = doc.add_paragraph(text)
    para.style.font.size = Pt(11)
    doc.save(file_obj)